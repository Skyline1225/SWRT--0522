from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import List, Dict, Any, Tuple
import re

class DocxHandler:
    def __init__(self):
        self.document = None
        self.paragraphs = []
        self.tables = []

    def read(self, file_path: str) -> bool:
        try:
            self.document = Document(file_path)
            self._parse_document()
            return True
        except Exception as e:
            print(f"读取Word文档失败: {e}")
            return False

    def _parse_document(self):
        self.paragraphs = []
        self.tables = []

        for para in self.document.paragraphs:
            if para.text.strip():
                self.paragraphs.append({
                    'text': para.text,
                    'style': para.style.name if para.style else 'Normal',
                    'alignment': str(para.alignment) if para.alignment else 'LEFT'
                })

        for table in self.document.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            self.tables.append(table_data)

    def get_all_text(self) -> str:
        text_parts = []
        for para in self.paragraphs:
            text_parts.append(para['text'])
        return '\n'.join(text_parts)

    def get_paragraphs(self) -> List[Dict[str, Any]]:
        return self.paragraphs

    def get_tables(self) -> List[List[List[str]]]:
        return self.tables

    def find_text(self, search_text: str) -> List[Tuple[int, str]]:
        matches = []
        for idx, para in enumerate(self.paragraphs):
            if search_text in para['text']:
                matches.append((idx, para['text']))
        return matches

    def replace_text(self, old_text: str, new_text: str) -> int:
        count = 0
        for para in self.document.paragraphs:
            if old_text in para.text:
                inline = para.runs
                for run in inline:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        count += 1
        return count

    def update_paragraph(self, index: int, new_text: str) -> bool:
        try:
            if index < len(self.document.paragraphs):
                para = self.document.paragraphs[index]
                para.clear()
                para.add_run(new_text)
                return True
        except Exception as e:
            print(f"更新段落失败: {e}")
        return False

    def update_table_cell(self, table_idx: int, row: int, col: int, new_value: str) -> bool:
        try:
            if table_idx < len(self.document.tables):
                table = self.document.tables[table_idx]
                if row < len(table.rows) and col < len(table.rows[row].cells):
                    cell = table.rows[row].cells[col]
                    cell.text = new_value
                    return True
        except Exception as e:
            print(f"更新表格单元格失败: {e}")
        return False

    def save(self, output_path: str) -> bool:
        try:
            if self.document:
                self.document.save(output_path)
                return True
        except Exception as e:
            print(f"保存Word文档失败: {e}")
        return False

    def add_paragraph(self, text: str, style: str = None) -> bool:
        try:
            if self.document:
                para = self.document.add_paragraph(text, style=style)
                return True
        except Exception as e:
            print(f"添加段落失败: {e}")
        return False

    def add_table(self, rows: int, cols: int, data: List[List[str]] = None) -> bool:
        try:
            if self.document:
                table = self.document.add_table(rows=rows, cols=cols)
                if data:
                    for i, row_data in enumerate(data):
                        if i < rows:
                            for j, cell_value in enumerate(row_data):
                                if j < cols:
                                    table.rows[i].cells[j].text = str(cell_value)
                return True
        except Exception as e:
            print(f"添加表格失败: {e}")
        return False
